"""DOCX / DOC Service – mask PII in Word documents."""
from fastapi import UploadFile
from typing import Dict, List
import docx
import io

from services.BaseService import BaseService
from utility.exceptions import FileValidationException, DocumentProcessingException


class DOCXService(BaseService):

    def _validate(self, document: UploadFile) -> None:
        fn = (document.filename or "").lower()
        if not (fn.endswith(".docx") or fn.endswith(".doc")):
            raise FileValidationException("File must be a .doc or .docx document")

    def _extract_text(self, raw: bytes, filename: str) -> str:
        try:
            doc = docx.Document(io.BytesIO(raw))
            parts = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    parts.extend(cell.text for cell in row.cells)
            return "\n".join(parts).strip()
        except Exception as e:
            raise DocumentProcessingException(f"DOCX extraction failed: {e}")

    def _build_masked_output(self, raw, mapping, anonymized_text, out_path):
        """Replace text in original DOCX with anonymized content, preserving styles."""
        try:
            doc = docx.Document(io.BytesIO(raw))
            anon_lines = anonymized_text.split("\n")
            idx = [0]  # mutable counter

            def _replace_para(para):
                if para.text.strip() and idx[0] < len(anon_lines):
                    if para.runs:
                        para.runs[0].text = anon_lines[idx[0]]
                        for r in para.runs[1:]:
                            r.text = ""
                    idx[0] += 1

            for para in doc.paragraphs:
                _replace_para(para)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            _replace_para(para)
            doc.save(out_path)
        except Exception as e:
            raise DocumentProcessingException(f"DOCX masking failed: {e}")
