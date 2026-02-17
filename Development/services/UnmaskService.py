"""
Unmask Service - de-anonymize documents by replacing tags with decrypted PII.
"""
from fastapi import UploadFile
from typing import Dict
import os
import tempfile
import logging

from repository.PIIRepository import PIIRepository
from utility.PresidioUtility import deanonymize_text
from utility.exceptions import DocumentProcessingException, DatabaseException

logger = logging.getLogger(__name__)


class UnmaskService:

    def __init__(self, repository: PIIRepository):
        self.repository = repository

    def process_document(self, request_id: str, document: UploadFile, input_type: str) -> Dict:
        temp_path = None
        out_path = None
        try:
            record = self.repository.get_pii_details(request_id)
            if not record:
                raise DatabaseException(f"No PII record for {request_id}")
            mapping = record.anonymizing_mapping
            key = record.encrypted_key
            if not mapping or not key:
                raise DatabaseException(f"Missing mapping/key for {request_id}")

            raw = document.file.read()
            
            # Use system temp directory
            temp_dir = tempfile.gettempdir()
            temp_path = os.path.join(temp_dir, f"{request_id}_masked.{input_type}")
            
            with open(temp_path, "wb") as f:
                f.write(raw)

            text = self._extract(raw, input_type)
            restored = deanonymize_text(text, mapping, key)

            out_path = os.path.join(temp_dir, f"{request_id}_unmasked.{input_type}")
            self._write_output(restored, out_path, input_type, raw)

            if temp_path and os.path.exists(temp_path):
                os.remove(temp_path)

            return {
                "request_id": request_id,
                "unmasked_document": out_path,
                "tags_replaced": len(mapping),
            }
        except Exception as e:
            for p in (temp_path, out_path):
                if p and os.path.exists(p):
                    os.remove(p)
            if isinstance(e, (DocumentProcessingException, DatabaseException)):
                raise
            raise DocumentProcessingException(f"De-anonymization failed: {e}")

    def _extract(self, raw: bytes, input_type: str) -> str:
        """Extract text from document bytes based on type."""
        try:
            if input_type in ("txt", "tavily"):
                try:
                    return raw.decode("utf-8")
                except UnicodeDecodeError:
                    return raw.decode("latin-1")
            elif input_type == "pdf":
                import PyPDF2, io
                reader = PyPDF2.PdfReader(io.BytesIO(raw))
                return "\n".join(p.extract_text() or "" for p in reader.pages).strip()
            elif input_type in ("docx", "doc"):
                import docx, io
                doc = docx.Document(io.BytesIO(raw))
                parts = [p.text for p in doc.paragraphs]
                for t in doc.tables:
                    for r in t.rows:
                        parts.extend(c.text for c in r.cells)
                return "\n".join(parts).strip()
            elif input_type == "csv":
                import pandas as pd, io
                df = pd.read_csv(io.BytesIO(raw))
                return df.to_string()
            elif input_type == "xlsx":
                import openpyxl, io
                wb = openpyxl.load_workbook(io.BytesIO(raw))
                parts = []
                for ws in wb.worksheets:
                    for row in ws.iter_rows():
                        parts.extend(str(c.value) for c in row if c.value)
                return " ".join(parts)
            elif input_type == "json":
                import json
                data = json.loads(raw.decode("utf-8"))
                return json.dumps(data, indent=2)
            else:
                return raw.decode("utf-8", errors="replace")
        except Exception as e:
            raise DocumentProcessingException(f"Text extraction failed: {e}")

    def _write_output(self, text: str, path: str, input_type: str, original_raw: bytes):
        """Write de-anonymized text back to the appropriate format."""
        try:
            if input_type in ("txt", "tavily", "json", "csv"):
                with open(path, "w", encoding="utf-8") as f:
                    f.write(text)
            elif input_type == "pdf":
                from reportlab.pdfgen import canvas
                from reportlab.lib.pagesizes import letter
                import io
                buf = io.BytesIO()
                c = canvas.Canvas(buf, pagesize=letter)
                y = 750
                for line in text.split("\n"):
                    if y < 50:
                        c.showPage()
                        y = 750
                    c.drawString(50, y, line[:95])
                    y -= 14
                c.save()
                with open(path, "wb") as f:
                    f.write(buf.getvalue())
            elif input_type in ("docx", "doc"):
                import docx as dx
                doc = dx.Document()
                for line in text.split("\n"):
                    doc.add_paragraph(line)
                doc.save(path)
            elif input_type == "xlsx":
                import openpyxl
                wb = openpyxl.Workbook()
                ws = wb.active
                for i, line in enumerate(text.split("\n"), 1):
                    ws.cell(row=i, column=1, value=line)
                wb.save(path)
            else:
                with open(path, "w", encoding="utf-8") as f:
                    f.write(text)
        except Exception as e:
            raise DocumentProcessingException(f"Output creation failed: {e}")
